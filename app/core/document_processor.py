"""
Multi-Format Document Processing Module
Handles text extraction from PDF, DOCX, TXT and other document formats
"""

import fitz  # PyMuPDF for PDF
from docx import Document as WordDocument  # python-docx for Word documents
import chardet  # For text encoding detection
import re
import os
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handles document processing operations for multiple file formats including
    PDF, DOCX, TXT with text extraction, cleaning, and intelligent chunking.
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF Document',
        '.docx': 'Microsoft Word Document',
        '.txt': 'Plain Text File'
    }
    
    def __init__(self, max_chunk_size: int = 4000):
        """
        Initialize document processor with configuration.
        
        Args:
            max_chunk_size: Maximum size for text chunks (in characters)
        """
        self.max_chunk_size = max_chunk_size
    
    def get_file_type(self, file_path: str) -> str:
        """
        Determine the file type from extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File extension in lowercase
        """
        return Path(file_path).suffix.lower()
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if the file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if format is supported
        """
        file_ext = self.get_file_type(file_path)
        return file_ext in self.SUPPORTED_FORMATS
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from supported document formats.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = self.get_file_type(file_path)
        
        if not self.is_supported_format(file_path):
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                return self._extract_docx_text(file_path)
            elif file_ext == '.txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """
        Extract text from PDF files using PyMuPDF.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            # Use explicit fitz.Document to avoid conflicts
            pdf_document = fitz.Document(pdf_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()
                
                # Clean the text
                page_text = self._clean_text(page_text)
                text += page_text + "\n\n"
            
            pdf_document.close()
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF: {pdf_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
            raise Exception(f"PDF processing failed: {str(e)}")
    
    def _extract_docx_text(self, docx_path: str) -> str:
        """
        Extract text from Word documents using python-docx.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = WordDocument(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Clean the text
            text = self._clean_text(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX: {docx_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    def _extract_txt_text(self, txt_path: str) -> str:
        """
        Extract text from plain text files with encoding detection.
        
        Args:
            txt_path: Path to the TXT file
            
        Returns:
            Extracted text content
        """
        try:
            # Detect encoding
            with open(txt_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
                if not encoding:
                    encoding = 'utf-8'  # fallback
            
            # Read with detected encoding
            with open(txt_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            # Clean the text
            text = self._clean_text(text)
            
            logger.info(f"Successfully extracted {len(text)} characters from TXT: {txt_path}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error processing TXT {txt_path}: {str(e)}")
            raise Exception(f"TXT processing failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single
        text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)  # Leading/trailing whitespace
        
        # Remove common OCR artifacts
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\@\#\$\%\^\&\*\+\=\<\>\~\`]', '', text)
        
        # Fix common formatting issues
        text = re.sub(r'(\w)(\d)', r'\1 \2', text)  # Add space between word and number
        text = re.sub(r'(\d)([A-Z])', r'\1 \2', text)  # Add space between number and capital letter
        
        return text
    
    def chunk_text(self, text: str, chunk_size: Optional[int] = None) -> List[str]:
        """
        Split text into manageable chunks for processing.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk (uses default if None)
            
        Returns:
            List of text chunks
        """
        if chunk_size is None:
            chunk_size = self.max_chunk_size
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        sentences = re.split(r'[.!?]+', text)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) + 1 > chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence
                else:
                    # Sentence is too long, split it by words
                    words = sentence.split()
                    while words:
                        word_chunk = ""
                        while words and len(word_chunk) + len(words[0]) + 1 <= chunk_size:
                            word_chunk += words.pop(0) + " "
                        if word_chunk:
                            chunks.append(word_chunk.strip())
            else:
                current_chunk += sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        logger.info(f"Text split into {len(chunks)} chunks")
        return chunks
    
    def get_document_info(self, file_path: str) -> Dict[str, any]:
        """
        Get basic information about the document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with document information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = self.get_file_type(file_path)
        file_size = os.path.getsize(file_path)
        
        info = {
            'filename': os.path.basename(file_path),
            'extension': file_ext,
            'format_name': self.SUPPORTED_FORMATS.get(file_ext, 'Unknown'),
            'file_size': file_size,
            'file_size_mb': round(file_size / (1024 * 1024), 2)
        }
        
        # Try to extract text and get character count
        try:
            text = self.extract_text(file_path)
            info.update({
                'character_count': len(text),
                'word_count': len(text.split()),
                'extraction_successful': True
            })
        except Exception as e:
            info.update({
                'character_count': 0,
                'word_count': 0,
                'extraction_successful': False,
                'error': str(e)
            })
        
        return info


# Maintain backward compatibility with old PDFProcessor
class PDFProcessor(DocumentProcessor):
    """
    Backward compatibility class - now inherits from DocumentProcessor
    """
    
    def __init__(self, max_chunk_size: int = 4000):
        super().__init__(max_chunk_size)
        logger.warning("PDFProcessor is deprecated. Use DocumentProcessor instead.")